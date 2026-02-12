#TFuerte/utilss/comprobante_word_generator
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
from pathlib import Path

def generar_comprobante_word(comprobante_data: dict, detalles_data: list, firmas_dir: str = None):
    """
    Genera un documento Word para un comprobante de salida de almacén.
    
    Args:
        comprobante_data (dict): Datos del comprobante.
        detalles_data (list): Lista de detalles del comprobante.
        firmas_dir (str): Directorio donde se encuentran las imágenes de las firmas.
    
    Returns:
        bytes: Contenido del archivo Word generado.
    """
    # Determinar la ruta de las firmas
    if firmas_dir is None:
        base_dir = Path(__file__).parent.parent.parent
        firmas_dir = str(base_dir / "assets" / "firmas")
        
        if not os.path.exists(firmas_dir):
            print(f"⚠️ No se encontró la carpeta de firmas en: {firmas_dir}")
            # Intentar ruta alternativa
            firmas_dir = str(base_dir / ".." / "assets" / "firmas")
            if not os.path.exists(firmas_dir):
                print(f"⚠️ Tampoco se encontró en: {firmas_dir}")
                firmas_dir = ""
            else:
                print(f"✅ Carpeta de firmas encontrada en: {firmas_dir}")
        else:
            print(f"✅ Carpeta de firmas encontrada en: {firmas_dir}")
    
    # Crear un nuevo documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título principal
    title = doc.add_heading('COMPROBANTE DE SALIDA DE ALMACÉN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del comprobante
    doc.add_heading('Información del Comprobante', level=1)
    
    # Crear tabla de información
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Datos del comprobante
    datos_info = [
        ("Destino:", comprobante_data.get("destino", "No especificado")),
        ("Fecha de Salida:", comprobante_data.get("fecha_salida", "No especificada")),
        ("Recibido por:", comprobante_data.get("recibido_por", "No especificado")),
        ("Observaciones:", comprobante_data.get("observaciones", "Ninguna"))
    ]
    
    for i, (label, value) in enumerate(datos_info):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    doc.add_paragraph()  # Espacio
    
    # Productos despachados
    doc.add_heading('Productos Despachados', level=1)
    
    if detalles_data:
        # Crear tabla de productos
        table_productos = doc.add_table(rows=len(detalles_data) + 1, cols=6)
        table_productos.style = 'Table Grid'
        table_productos.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        headers = ["No.", "Código", "Descripción", "U/M", "Cantidad", "Importe"]
        header_cells = table_productos.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar con datos de productos
        total_general = 0
        for i, detalle in enumerate(detalles_data, start=1):
            row = table_productos.rows[i]
            
            # Número de orden
            row.cells[0].text = str(i)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Código
            row.cells[1].text = detalle.get("codigo", "")
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Descripción
            row.cells[2].text = detalle.get("descripcion", "")
            
            # Unidad de medida
            row.cells[3].text = detalle.get("um", "")
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cantidad
            row.cells[4].text = str(detalle.get("cantidad", 0))
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Importe
            importe = detalle.get("importe", 0)
            row.cells[5].text = f"${importe:.2f}"
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            total_general += importe
        
        # Agregar fila de total
        doc.add_paragraph()  # Espacio
        total_para = doc.add_paragraph()
        total_para.add_run(f"TOTAL: ${total_general:.2f}").bold = True
        total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Firmas - CON CARGOS
    doc.add_heading('Firmas de Autorización', level=1)
    
    # Tabla de firmas con 3 columnas
    firmas_table = doc.add_table(rows=4, cols=3)  # 4 filas: encabezado, nombre, cargo, firma
    firmas_table.style = 'Table Grid'
    firmas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    headers_firmas = ["Aprobado por", "Despachado por", "Recibido por"]
    header_cells = firmas_table.rows[0].cells
    for i, header in enumerate(headers_firmas):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Nombres (fila 1)
    nombres = [
        comprobante_data.get("aprobado_por", ""),
        comprobante_data.get("despachado_por", ""),
        comprobante_data.get("recibido_por", "")
    ]
    
    for i, nombre in enumerate(nombres):
        firmas_table.rows[1].cells[i].text = nombre
        firmas_table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cargos (fila 2)
    cargos = [
        comprobante_data.get("cargo_aprobado", "Administrador"),
        comprobante_data.get("cargo_despachado", "Jefe de Almacén"),
        "Solicitante"
    ]
    
    for i, cargo in enumerate(cargos):
        firmas_table.rows[2].cells[i].text = cargo
        firmas_table.rows[2].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        firmas_table.rows[2].cells[i].paragraphs[0].runs[0].italic = True
    
    # Intentar agregar imágenes de firmas (fila 3)
    try:
        firmas_map = {
            comprobante_data.get("aprobado_por", ""): comprobante_data.get("firma_aprobado", ""),
            comprobante_data.get("despachado_por", ""): comprobante_data.get("firma_despachado", ""),
            comprobante_data.get("recibido_por", ""): comprobante_data.get("firma_recibido", "")
        }
        
        for i, (nombre, firma_archivo) in enumerate(firmas_map.items()):
            if nombre and firma_archivo and os.path.exists(os.path.join(firmas_dir, firma_archivo)):
                cell = firmas_table.rows[3].cells[i]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run()
                run.add_picture(os.path.join(firmas_dir, firma_archivo), width=Inches(1.2))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Si no hay firma, poner línea para firmar
                firmas_table.rows[3].cells[i].text = "___________________"
                firmas_table.rows[3].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                firmas_table.rows[3].cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    except Exception as e:
        print(f"⚠️ Error agregando firmas: {e}")
        import traceback
        traceback.print_exc()
        for i in range(3):
            firmas_table.rows[3].cells[i].text = "___________________"
            firmas_table.rows[3].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha de generación
    doc.add_paragraph()
    fecha_gen = doc.add_paragraph(f"Documento generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    fecha_gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Guardar en memoria
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()