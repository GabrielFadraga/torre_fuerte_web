# TFuerte/utilss/word_generator_fin.py - VERSIÓN CORREGIDA
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any

def generar_word_solicitud_fin(numero_solicitud: str, recursos: List[Dict[str, Any]], solicitante: Dict[str, Any], firmas_dir: str = None):
    """
    Genera documento Word para solicitud de financiamiento.
    NOTA: No incluye área logística, solo financiero y administración.
    """
    if not recursos:
        print(f"❌ No se encontraron recursos para la solicitud {numero_solicitud}")
        return None
    
    primer_recurso = recursos[0]
    
    # Determinar la ruta de las firmas
    if firmas_dir is None:
        base_dir = Path(__file__).parent.parent.parent
        firmas_dir = str(base_dir / "assets" / "firmas")
        
        if not os.path.exists(firmas_dir):
            firmas_dir = str(base_dir / ".." / "assets" / "firmas")
            if not os.path.exists(firmas_dir):
                firmas_dir = ""
    
    # Crear un nuevo documento
    doc = Document()
    
    # Configurar márgenes ajustados
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título principal
    title = doc.add_heading('SOLICITUD DE FINANCIAMIENTO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    
    # INFORMACIÓN GENERAL DE LA SOLICITUD
    title_solicitud = doc.add_heading('Datos de la Solicitud', level=1)
    title_solicitud.paragraph_format.space_after = Pt(0)
    
    # Obtener datos del solicitante
    solicitante_usuario = solicitante.get("usuario", "Solicitante")
    solicitante_cargo = solicitante.get("cargo", "Solicitante")
    
    # Información en tabla compacta
    datos_table = doc.add_table(rows=7, cols=2)
    datos_table.style = 'Table Grid'
    
    datos = [
        ("Área solicitante:", primer_recurso.get("Area solicitante", "No especificado")),
        ("Fecha:", primer_recurso.get("Fecha", "No especificada")),
        ("Número contrato/suplemento:", primer_recurso.get("Numero de contrato/suplemento", "No especificado")),
        ("Orden de trabajo:", primer_recurso.get("Orden de trabajo", "No especificada")),
        ("Número de solicitud:", numero_solicitud),
        ("Solicitante:", solicitante_usuario),
        ("Cargo del solicitante:", solicitante_cargo),
    ]
    
    for i, (label, value) in enumerate(datos):
        row = datos_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
        row.cells[1].text = value
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    
    # Tabla de productos
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    title_productos = doc.add_heading('Productos Solicitados', level=1)
    title_productos.paragraph_format.space_after = Pt(0)
    
    if recursos:
        # Crear tabla de productos
        table_productos = doc.add_table(rows=len(recursos) + 1, cols=6)
        table_productos.style = 'Table Grid'
        table_productos.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        headers = ["No.", "Servicio/Tipo", "Descripción", "Cantidad", "Precio Unitario", "Importe"]
        header_cells = table_productos.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar con datos de recursos
        total_general = 0
        for i, recurso in enumerate(recursos, start=1):
            row = table_productos.rows[i]
            
            # Número de orden
            row.cells[0].text = str(i)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Servicio/Tipo
            row.cells[1].text = recurso.get("Servicio", "")
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Descripción
            row.cells[2].text = recurso.get("Descripcion", "")
            
            # Cantidad
            row.cells[3].text = str(recurso.get("Cantidad", ""))
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Precio unitario (precio base + 25%)
            precio_unitario = recurso.get("Precio unitario", 0)
            row.cells[4].text = f"${precio_unitario:.2f}"
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Importe
            importe = recurso.get("Importe", 0)
            row.cells[5].text = f"${importe:.2f}"
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            total_general += importe
    
    # Total general
    total_para = doc.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_para.paragraph_format.space_before = Pt(6)
    total_para.paragraph_format.space_after = Pt(0)
    total_run = total_para.add_run(f"TOTAL GENERAL: ${total_general:.2f}")
    total_run.bold = True
    total_run.font.size = Pt(12)
    
    # NOTA IMPORTANTE: Para financiamiento, SOLO van:
    # 1. Solicitante
    # 2. Revisor Financiero (Alexander)
    # 3. Administrador/Presidente (Maikel)
    # NO incluye logística
    
    # Aprobaciones y firmas
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    title_aprobaciones = doc.add_heading('Aprobaciones y Firmas', level=1)
    title_aprobaciones.paragraph_format.space_after = Pt(0)
    
    # Tabla de aprobaciones (SOLO 4 filas incluyendo encabezado)
    aprobaciones_table = doc.add_table(rows=4, cols=4)  # CAMBIADO de 5 a 4 filas
    aprobaciones_table.style = 'Table Grid'
    
    # Encabezados
    headers = ["Rol", "Nombre", "Cargo", "Firma"]
    header_cells = aprobaciones_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Datos de aprobaciones - SOLO FINANCIERO Y ADMIN
    aprobaciones_data = [
        ("Solicitado por:", solicitante_usuario, solicitante_cargo, ""),
        ("Revisado por:", "Alexander", "Revisor Financiero", ""),
        ("Aprobado por:", "Maikel", "Presidente", ""),
        # NOTA: NO incluir "Recibido por:" para financiamiento
    ]
    
    for i, (rol, nombre, cargo, firma) in enumerate(aprobaciones_data):
        row = aprobaciones_table.rows[i + 1]
        row.cells[0].text = rol
        row.cells[1].text = nombre
        row.cells[2].text = cargo
    
    # Intentar agregar imágenes de firmas
    try:
        # Mapeo de nombres a archivos de firmas
        firmas_map = {
            "Alexander": "Alexander.png",
            "Maikel": "maykeltf.png",
            # NOTA: Miguel NO va en financiamiento
        }
        
        # Buscar firma del solicitante
        firma_solicitante = None
        if solicitante_usuario:
            posibles_nombres = [
                f"{solicitante_usuario}.png",
                f"{solicitante_usuario}.jpg",
                f"{solicitante_usuario}.jpeg",
                f"{solicitante_usuario.lower()}.png",
                f"{solicitante_usuario.capitalize()}.png",
                f"{solicitante_usuario.split()[0]}.png" if ' ' in solicitante_usuario else None,
            ]
            
            for nombre_archivo in posibles_nombres:
                if nombre_archivo and nombre_archivo != "None.png":
                    ruta_prueba = os.path.join(firmas_dir, nombre_archivo)
                    if os.path.exists(ruta_prueba):
                        firma_solicitante = nombre_archivo
                        break
        
        # Agregar firmas a las celdas (solo 3 filas de datos ahora)
        for i in range(1, 4):  # CAMBIADO: de range(1, 5) a range(1, 4)
            row = aprobaciones_table.rows[i]
            nombre_celda = row.cells[1].text.strip()
            
            firma_file = None
            
            if i == 1 and firma_solicitante:
                # Firma del solicitante
                firma_file = os.path.join(firmas_dir, firma_solicitante)
            elif nombre_celda in firmas_map:
                # Firma de los otros roles
                nombre_archivo = firmas_map[nombre_celda]
                firma_file = os.path.join(firmas_dir, nombre_archivo)
            
            if firma_file and os.path.exists(firma_file):
                try:
                    row.cells[3].text = ""
                    para = row.cells[3].paragraphs[0]
                    run = para.add_run()
                    run.add_picture(firma_file, width=Inches(1.0))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as img_error:
                    print(f"❌ Error al agregar imagen: {img_error}")
                    row.cells[3].text = "Firma"
                    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row.cells[3].text = "Firma"
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    except Exception as e:
        print(f"⚠️ Error general agregando firmas: {e}")
        import traceback
        traceback.print_exc()
        for i in range(1, 4):  # CAMBIADO: de range(1, 5) a range(1, 4)
            aprobaciones_table.rows[i].cells[3].text = "Firma"
            aprobaciones_table.rows[i].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha de generación
    fecha_gen = doc.add_paragraph(f"Documento generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    fecha_gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha_gen.paragraph_format.space_before = Pt(6)
    fecha_gen.paragraph_format.space_after = Pt(0)
    
    # Nota sobre el cálculo
    nota = doc.add_paragraph("Nota: El precio unitario incluye un 25% adicional al precio base del producto.")
    nota.paragraph_format.space_before = Pt(3)
    nota.paragraph_format.space_after = Pt(0)
    nota.runs[0].italic = True
    nota.runs[0].font.size = Pt(8)
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Guardar en memoria
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()