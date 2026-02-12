from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from pathlib import Path

def set_column_width(column, width):
    """Establece el ancho de una columna en centímetros usando XML"""
    for cell in column.cells:
        cell.width = width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.first_child_found_in("w:tcW")
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            tcPr.append(tcW)
        tcW.set(qn('w:w'), str(width))
        tcW.set(qn('w:type'), 'dxa')  # Tipo de medida: dxa = 1/1440 de pulgada

def set_row_height(row, height):
    """Establece la altura de una fila usando XML"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), "atLeast")  # "atLeast" o "exact"
    trPr.append(trHeight)

def generar_word_solicitud_rm(solicitud_rm: dict, solicitante: dict, firmas_dir: str = None):
    """
    Genera un documento Word para una solicitud de recursos y materiales.
    Ahora maneja múltiples recursos.
    
    Args:
        solicitud_rm (dict): Datos de la solicitud (debe incluir campo 'recursos').
        solicitante (dict): Datos del solicitante.
        firmas_dir (str): Directorio donde se encuentran las imágenes de las firmas.
    
    Returns:
        bytes: Contenido del archivo Word generado.
    """
    # Determinar la ruta de las firmas
    if firmas_dir is None:
        base_dir = Path(__file__).parent.parent.parent
        firmas_dir = str(base_dir / "assets" / "firmas")
        
        if not os.path.exists(firmas_dir):
            print(f"⚠️ No se encontró la carpeta de firmas en: {firmas_dir}")
            firmas_dir = str(base_dir / ".." / "assets" / "firmas")
            if not os.path.exists(firmas_dir):
                print(f"⚠️ Tampoco se encontró en: {firmas_dir}")
                firmas_dir = ""
            else:
                print(f"✅ Carpeta de firmas encontrada en: {firmas_dir}")
        else:
            print(f"✅ Carpeta de firmas encontrada en: {firmas_dir}")
    
    # Crear un nuevo documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título principal - SIN ESPACIO DESPUÉS
    title = doc.add_heading('SOLICITUD DE RECURSOS Y MATERIALES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)  # Sin espacio después del título
    
    # INFORMACIÓN GENERAL DE LA SOLICITUD - SIN TABLA
    # Título SIN espacio después
    title_solicitud = doc.add_heading('Datos de la Solicitud', level=1)
    title_solicitud.paragraph_format.space_after = Pt(0)  # Sin espacio
    
    # Crear un formato más profesional sin tabla
    # Agregamos un párrafo con el centro de costo
    centro_costo_para = doc.add_paragraph()
    centro_costo_para.paragraph_format.space_after = Pt(0)  # Sin espacio
    centro_costo_run = centro_costo_para.add_run('Centro de costo: ')
    centro_costo_run.bold = True
    centro_costo_para.add_run(solicitud_rm.get("Centro costo", "No especificado"))
    
    # Agregar un párrafo con la fecha
    fecha_para = doc.add_paragraph()
    fecha_para.paragraph_format.space_after = Pt(0)  # Sin espacio
    fecha_run = fecha_para.add_run('Fecha: ')
    fecha_run.bold = True
    fecha_para.add_run(solicitud_rm.get("Fecha", "No especificada"))
    
    # Agregar un párrafo con la orden de trabajo
    orden_para = doc.add_paragraph()
    orden_para.paragraph_format.space_after = Pt(0)  # Sin espacio
    orden_run = orden_para.add_run('Orden de trabajo: ')
    orden_run.bold = True
    orden_para.add_run(solicitud_rm.get("Orden trabajo", "No especificada"))
    
    # Agregar un párrafo con el solicitante - CON UN POCO DE ESPACIO
    solicitante_para = doc.add_paragraph()
    solicitante_para.paragraph_format.space_after = Pt(6)  # Poco espacio antes de la tabla
    solicitante_run = solicitante_para.add_run('Solicitante: ')
    solicitante_run.bold = True
    solicitante_para.add_run(solicitante.get("usuario", "No especificado"))
    
    # TABLA DE RECURSOS (NUEVO)
    # Título SIN espacio después
    title_recursos = doc.add_heading('Recursos Solicitados', level=1)
    title_recursos.paragraph_format.space_after = Pt(0)  # Sin espacio
    
    # Obtener recursos
    recursos = solicitud_rm.get("recursos", [])
    
    # CALCULAR ANCHO TOTAL DE LA TABLA DE RECURSOS
    # Este será el ancho que también usaremos para la tabla de aprobaciones
    # Anchura en twips para tabla de recursos:
    # Total: 1.5 + 8 + 2 + 2 + 5 = 18.5 cm aprox * 567 = ~10490 twips
    
    if recursos:
        # Crear tabla de recursos (fila de encabezado + filas de recursos)
        table_recursos = doc.add_table(rows=len(recursos) + 1, cols=5)
        table_recursos.style = 'Table Grid'
        table_recursos.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # CONFIGURACIÓN DE ANCHOS DE COLUMNAS MODIFICADA - USANDO XML DIRECTAMENTE
        # No. - columna más estrecha (1.5 cm)
        # Descripción - columna más ancha (8 cm)
        # U/M - columna estrecha (2 cm)
        # Cantidad - columna estrecha (2 cm)
        # Observaciones - columna ancha (5 cm)
        
        # Convertir a medidas en twips (1 cm = 567 twips)
        widths_twips = [
            850,   # No. - 1.5 cm aprox
            4536,  # Descripción - 8 cm aprox
            1134,  # U/M - 2 cm aprox
            1134,  # Cantidad - 2 cm aprox
            2835   # Observaciones - 5 cm aprox
        ]
        
        # Calcular el ancho total
        total_width_recursos = sum(widths_twips)
        
        # Configurar anchos usando XML
        for i, width in enumerate(widths_twips):
            set_column_width(table_recursos.columns[i], width)
        
        # Encabezados
        headers = ["No.", "Descripción", "U/M", "Cantidad", "Observaciones"]
        header_cells = table_recursos.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Llenar con datos de recursos
        for i, recurso in enumerate(recursos, start=1):
            row = table_recursos.rows[i]
            
            # Número de orden
            row.cells[0].text = str(i)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Descripción
            row.cells[1].text = recurso.get("descripcion", "")
            
            # Unidad de medida
            row.cells[2].text = recurso.get("unidad_medida", "")
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cantidad
            row.cells[3].text = str(recurso.get("cantidad", ""))
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Observaciones
            row.cells[4].text = recurso.get("observaciones", "")
            
            # Ajustar altura de fila si es necesario
            if len(recurso.get("descripcion", "")) > 100 or len(recurso.get("observaciones", "")) > 100:
                set_row_height(row, 500)  # Fila más alta para contenido largo
    else:
        no_recursos = doc.add_paragraph("No hay recursos registrados para esta solicitud.")
        no_recursos.paragraph_format.space_after = Pt(6)
    
    # Espacio pequeño antes de la siguiente sección
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    # Aprobaciones y firmas
    # Título SIN espacio después
    title_aprobaciones = doc.add_heading('Aprobaciones y Firmas', level=1)
    title_aprobaciones.paragraph_format.space_after = Pt(0)  # Sin espacio
    
    # Tabla de aprobaciones - Mismo ancho total que la tabla de recursos
    aprobaciones_table = doc.add_table(rows=5, cols=4)
    aprobaciones_table.style = 'Table Grid'
    aprobaciones_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Distribuir el ancho total entre las 4 columnas de aprobaciones
    # Manteniendo proporciones apropiadas
    # Usamos el mismo ancho total: ~10490 twips
    
    # Distribución proporcional para 4 columnas:
    # Rol: 20%, Nombre: 25%, Cargo: 25%, Firma: 30%
    widths_aprobaciones_twips = [
        int(total_width_recursos * 0.20),  # Rol - 20%
        int(total_width_recursos * 0.25),  # Nombre - 25%
        int(total_width_recursos * 0.25),  # Cargo - 25%
        int(total_width_recursos * 0.30),  # Firma - 30%
    ]
    
    # Ajustar para que la suma sea exactamente el total
    diff = total_width_recursos - sum(widths_aprobaciones_twips)
    if diff != 0:
        widths_aprobaciones_twips[3] += diff  # Ajustar en la última columna
    
    for i, width in enumerate(widths_aprobaciones_twips):
        set_column_width(aprobaciones_table.columns[i], width)
    
    # Encabezados
    headers = ["Rol", "Nombre", "Cargo", "Firma"]
    header_cells = aprobaciones_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Datos de aprobaciones
    aprobaciones_data = [
        ("Solicitado por:", solicitante.get("usuario", ""), solicitante.get("cargo", "Solicitante"), ""),
        ("Revisado por:", "Alexander", "Jefe de Área Técnica", ""),
        ("Aprobado por:", "Maikel", "Presidente", ""),
        ("Recibido por:", "Miguel", "Jefe Área Logística", "")
    ]
    
    for i, (rol, nombre, cargo, firma) in enumerate(aprobaciones_data):
        row = aprobaciones_table.rows[i + 1]
        row.cells[0].text = rol
        row.cells[1].text = nombre
        row.cells[2].text = cargo
        
        # Altura de fila un poco más alta para que las firmas quepan bien
        set_row_height(row, 600)  # 600 twips = aproximadamente 1.1 cm
    
    # Intentar agregar imágenes de firmas
    try:
        firmas_map = {
            "Alexander": "Alexander.png",
            "Maikel": "maykeltf.png",
            "Miguel": "Miguel.png",
        }
        
        solicitante_usuario = solicitante.get("usuario", "").strip()
        print(f"🔍 Buscando firma para solicitante: '{solicitante_usuario}'")
        
        # Buscar firma del solicitante
        firma_solicitante = None
        if solicitante_usuario:
            posibles_nombres = [
                f"{solicitante_usuario}.png",
                f"{solicitante_usuario}.jpg",
                f"{solicitante_usuario}.jpeg",
                f"{solicitante_usuario.lower()}.png",
                f"{solicitante_usuario.capitalize()}.png",
            ]
            
            for nombre_archivo in posibles_nombres:
                ruta_prueba = os.path.join(firmas_dir, nombre_archivo)
                if os.path.exists(ruta_prueba):
                    firma_solicitante = nombre_archivo
                    print(f"✅ Encontrada firma del solicitante: {firma_solicitante}")
                    break
        
        # Agregar firmas a las celdas
        for i in range(1, 5):
            row = aprobaciones_table.rows[i]
            nombre_celda = row.cells[1].text.strip()
            
            firma_file = None
            
            if i == 1 and firma_solicitante:
                firma_file = os.path.join(firmas_dir, firma_solicitante)
            elif nombre_celda in firmas_map:
                nombre_archivo = firmas_map[nombre_celda]
                firma_file = os.path.join(firmas_dir, nombre_archivo)
            
            if firma_file and os.path.exists(firma_file):
                try:
                    row.cells[3].text = ""
                    para = row.cells[3].paragraphs[0]
                    run = para.add_run()
                    # Tamaño de imagen ajustado al nuevo ancho de columna
                    run.add_picture(firma_file, width=Inches(1.0))  # Aumentado a 1.0 pulgadas
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as img_error:
                    print(f"❌ Error al agregar imagen: {img_error}")
                    row.cells[3].text = "Firma"
                    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row.cells[3].text = "Firma"
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    except Exception as e:
        print(f"⚠️ Error general agregando firmas: {e}")
        import traceback
        traceback.print_exc()
        for i in range(1, 5):
            aprobaciones_table.rows[i].cells[3].text = "Firma"
            aprobaciones_table.rows[i].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Espacio pequeño antes de la fecha
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    # Fecha de generación
    fecha_gen = doc.add_paragraph(f"Documento generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    fecha_gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fecha_gen.paragraph_format.space_before = Pt(12)
    
    # Guardar en memoria
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()